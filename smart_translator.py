#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Smart Translator - 基于大模型的智能翻译技能
支持网页、Word、PDF、Markdown 文档翻译，智能分批处理，断点续翻
"""

import os
import sys
import json
import hashlib
import asyncio
import aiohttp
import time
from pathlib import Path
from typing import List, Dict, Optional
from datetime import datetime

# 文档解析依赖
try:
    from docx import Document
    import pdfplumber
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False
    print("警告：文档处理库未安装，请运行：pip install python-docx pdfplumber reportlab")


class TranslationCheckpoint:
    """翻译进度管理 - 支持断点续翻"""
    
    def __init__(self, doc_path: str):
        self.doc_id = hashlib.md5(doc_path.encode()).hexdigest()
        self.cache_dir = Path(".translation_cache")
        self.cache_dir.mkdir(exist_ok=True)
        self.state_file = self.cache_dir / f"{self.doc_id}.json"
    
    def save_batch(self, batch_id: int, result: str, cost: float = 0):
        """保存批次结果"""
        state = self.load()
        state['completed'].append(batch_id)
        state['results'][str(batch_id)] = {
            'text': result,
            'cost': cost,
            'timestamp': time.time()
        }
        state['updated_at'] = time.time()
        
        with open(self.state_file, 'w', encoding='utf-8') as f:
            json.dump(state, f, ensure_ascii=False, indent=2)
    
    def is_completed(self, batch_id: int) -> bool:
        """检查批次是否已完成"""
        state = self.load()
        return batch_id in state['completed']
    
    def get_result(self, batch_id: int) -> Optional[str]:
        """获取已完成的批次结果"""
        state = self.load()
        data = state['results'].get(str(batch_id))
        return data['text'] if data else None
    
    def load(self) -> Dict:
        """加载状态"""
        if self.state_file.exists():
            with open(self.state_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {
            'completed': [],
            'results': {},
            'created_at': time.time(),
            'total_cost': 0
        }
    
    def clear(self):
        """清除缓存"""
        if self.state_file.exists():
            self.state_file.unlink()
    
    def get_progress(self, total: int) -> Dict:
        """获取进度信息"""
        state = self.load()
        completed = len(state['completed'])
        return {
            'total': total,
            'completed': completed,
            'remaining': total - completed,
            'progress': completed / total * 100,
            'cost': state.get('total_cost', 0)
        }


class SmartTranslator:
    """智能翻译官 - 基于大模型 API"""
    
    def __init__(self, model: str = "qwen-plus", mode: str = "strict", api_key: Optional[str] = None):
        self.model = model
        self.mode = mode
        self.api_key = api_key or os.getenv("DASHSCOPE_API_KEY")
        self.cache_dir = Path(".translation_cache")
        self.cache_dir.mkdir(exist_ok=True)
        
        # 模型配置
        self.model_config = {
            "qwen-turbo": {"max_tokens": 32000, "batch_size": 15000, "price": 0.002},
            "qwen-plus": {"max_tokens": 32000, "batch_size": 10000, "price": 0.01},
            "qwen-max": {"max_tokens": 32000, "batch_size": 8000, "price": 0.02},
            "gpt-4o": {"max_tokens": 128000, "batch_size": 50000, "price": 0.03},
            "claude-3.5-sonnet": {"max_tokens": 200000, "batch_size": 80000, "price": 0.05},
        }
    
    def build_prompt(self, text: str, mode: str = "strict") -> str:
        """构建翻译 Prompt"""
        prompts = {
            "strict": f"""请翻译以下内容为中文。
要求：
- 100% 中文输出，不允许英文混杂
- 保留通用缩写（AI, API, UI, UX, SDK 等）
- 技术品牌名保留英文（React, Docker, Kubernetes, GitHub 等）
- 保持原文格式和结构

内容：
{text}""",
            
            "technical": f"""请翻译以下技术文档为中文。
要求：
- 保留通用缩写（API, UI, UX, AI, SDK, HTTP, JSON 等）
- 技术品牌名保留英文（React, Vue, Angular, Docker, Kubernetes, MySQL 等）
- 专业术语准确翻译
- 保持代码和技术名词不变

内容：
{text}""",
            
            "fluent": f"""请翻译以下内容为流畅的中文。
要求：
- 优化可读性和流畅度
- 适当调整句式结构符合中文习惯
- 保留必要的专业术语和缩写
- 让译文自然易懂

内容：
{text}""",
        }
        
        return prompts.get(mode, prompts["strict"])
    
    async def call_llm(self, prompt: str) -> str:
        """调用大模型 API"""
        if not self.api_key:
            raise ValueError("API Key 未配置，请设置 DASHSCOPE_API_KEY 环境变量")
        
        # 阿里云百炼 API
        url = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": self.model,
            "input": {
                "messages": [
                    {"role": "system", "content": "你是一个专业的翻译助手。"},
                    {"role": "user", "content": prompt}
                ]
            },
            "parameters": {
                "temperature": 0.3,
                "max_tokens": 4000
            }
        }
        
        async with aiohttp.ClientSession() as session:
            async with session.post(url, headers=headers, json=payload) as response:
                if response.status == 200:
                    data = await response.json()
                    return data['output']['choices'][0]['message']['content']
                else:
                    error = await response.text()
                    raise Exception(f"API 调用失败：{error}")
    
    def smart_split(self, content: str, max_batch_size: int = 10000) -> List[str]:
        """智能分批 - 按语义分段"""
        # 按章节分割（优先）
        import re
        chapters = re.split(r'\n(?=#{1,3} )', content)
        
        batches = []
        current_batch = []
        current_length = 0
        
        for chapter in chapters:
            chapter_length = len(chapter)
            
            if chapter_length > max_batch_size:
                # 大章节按段落分割
                paragraphs = chapter.split('\n\n')
                for para in paragraphs:
                    if current_length + len(para) > max_batch_size:
                        batches.append('\n\n'.join(current_batch))
                        current_batch = []
                        current_length = 0
                    current_batch.append(para)
                    current_length += len(para)
            else:
                # 整章作为一批
                batches.append(chapter)
        
        if current_batch:
            batches.append('\n\n'.join(current_batch))
        
        return batches
    
    def parse_document(self, file_path: str) -> tuple:
        """解析文档，提取文本和结构"""
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext == '.docx':
            return self._parse_docx(path)
        elif ext == '.md':
            return self._parse_markdown(path)
        elif ext == '.txt':
            return self._parse_txt(path)
        elif ext == '.pdf':
            return self._parse_pdf(path)
        else:
            raise ValueError(f"不支持的文件格式：{ext}")
    
    def _parse_docx(self, path: Path) -> tuple:
        """解析 Word 文档"""
        if not DOC_SUPPORT:
            raise ImportError("python-docx 未安装")
        
        doc = Document(path)
        content = []
        structure = {'paragraphs': []}
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                content.append(text)
                structure['paragraphs'].append({
                    'index': i,
                    'text': text,
                    'style': para.style.name
                })
        
        return '\n\n'.join(content), structure
    
    def _parse_markdown(self, path: Path) -> tuple:
        """解析 Markdown 文件"""
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        structure = {'headers': [], 'code_blocks': []}
        # 简单提取结构
        import re
        for match in re.finditer(r'^(#{1,6})\s+(.+)$', content, re.MULTILINE):
            structure['headers'].append({
                'level': len(match.group(1)),
                'text': match.group(2)
            })
        
        return content, structure
    
    def _parse_txt(self, path: Path) -> tuple:
        """解析 TXT 文件"""
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content, {'type': 'plain_text'}
    
    def _parse_pdf(self, path: Path) -> tuple:
        """解析 PDF 文件"""
        if not DOC_SUPPORT:
            raise ImportError("pdfplumber 未安装")
        
        content = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    content.append(text)
        
        return '\n\n'.join(content), {'pages': len(content)}
    
    async def translate_batch(self, batch_text: str, checkpoint: TranslationCheckpoint, batch_id: int) -> str:
        """翻译单个批次"""
        # 检查是否已完成
        if checkpoint.is_completed(batch_id):
            print(f"  批次 {batch_id} 已完成，跳过")
            return checkpoint.get_result(batch_id)
        
        # 构建 Prompt
        prompt = self.build_prompt(batch_text, self.mode)
        
        # 调用大模型
        print(f"  翻译批次 {batch_id}...")
        result = await self.call_llm(prompt)
        
        # 估算成本
        tokens = len(batch_text) // 4  # 粗略估算
        cost = tokens * self.model_config.get(self.model, {}).get('price', 0.01) / 1000
        
        # 保存进度
        checkpoint.save_batch(batch_id, result, cost)
        
        return result
    
    async def translate_batches(self, batches: List[str], checkpoint: TranslationCheckpoint, max_concurrent: int = 5) -> List[str]:
        """并发翻译批次"""
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def translate_with_semaphore(batch_id: int, batch_text: str):
            async with semaphore:
                return await self.translate_batch(batch_text, checkpoint, batch_id)
        
        # 创建任务
        tasks = [
            translate_with_semaphore(i, batch)
            for i, batch in enumerate(batches)
        ]
        
        results = []
        total = len(batches)
        
        for coro in asyncio.as_completed(tasks):
            result = await coro
            results.append(result)
            
            # 更新进度显示
            completed = len(results)
            progress = checkpoint.get_progress(total)
            print(f"  进度：{completed}/{total} ({progress['progress']:.1f}%) - 成本：¥{progress['cost']:.2f}")
        
        return results
    
    def save_document(self, translated: str, structure: Dict, original_path: str, output_path: str):
        """保存翻译后的文档"""
        path = Path(original_path)
        ext = path.suffix.lower()
        
        if ext == '.docx':
            self._save_docx(translated, structure, output_path)
        elif ext in ['.md', '.txt']:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(translated)
        elif ext == '.pdf':
            self._save_pdf(translated, output_path)
        else:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(translated)
    
    def _save_docx(self, translated: str, structure: Dict, output_path: str):
        """保存为 Word 文档"""
        if not DOC_SUPPORT:
            raise ImportError("python-docx 未安装")
        
        doc = Document()
        paragraphs = translated.split('\n\n')
        
        for text in paragraphs:
            if text.strip():
                doc.add_paragraph(text)
        
        doc.save(output_path)
    
    def _save_pdf(self, translated: str, output_path: str):
        """保存为 PDF"""
        if not DOC_SUPPORT:
            raise ImportError("reportlab 未安装")
        
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        
        y = height - 50
        lines = translated.split('\n')
        
        for line in lines:
            if y < 50:
                c.showPage()
                y = height - 50
            
            # 简单处理长行
            max_chars = 80
            while len(line) > max_chars:
                split_pos = line.rfind(' ', 0, max_chars)
                if split_pos == -1:
                    split_pos = max_chars
                c.drawString(50, y, line[:split_pos])
                y -= 15
                line = line[split_pos:].lstrip()
            
            c.drawString(50, y, line)
            y -= 15
        
        c.save()
    
    async def translate_document(self, file_path: str, output_path: Optional[str] = None, max_concurrent: int = 5) -> str:
        """翻译文档"""
        print(f"\n开始翻译：{file_path}")
        print("=" * 60)
        
        # 1. 解析文档
        print("1. 解析文档...")
        content, structure = self.parse_document(file_path)
        total_chars = len(content)
        print(f"   总字数：{total_chars:,} 字")
        
        # 2. 智能分批
        print("2. 智能分批...")
        config = self.model_config.get(self.model, {})
        batch_size = config.get('batch_size', 10000)
        batches = self.smart_split(content, batch_size)
        print(f"   分批数：{len(batches)} 批 (每批约{batch_size:,}字)")
        
        # 3. 创建检查点
        print("3. 创建进度检查点...")
        checkpoint = TranslationCheckpoint(file_path)
        
        # 检查是否有未完成的翻译
        progress = checkpoint.get_progress(len(batches))
        if progress['completed'] > 0:
            print(f"   发现未完成的翻译：已完成 {progress['completed']}/{len(batches)}")
            print(f"   已用成本：¥{progress['cost']:.2f}")
            confirm = input("   是否从断点处继续？[Y/n]: ").strip().lower()
            if confirm == 'n':
                checkpoint.clear()
                print("   已清除缓存，重新开始")
        
        # 4. 并发翻译
        print("4. 开始翻译...")
        start_time = time.time()
        results = await self.translate_batches(batches, checkpoint, max_concurrent)
        elapsed = time.time() - start_time
        
        # 5. 合并结果
        print("5. 合并结果...")
        final = '\n\n'.join(results)
        
        # 6. 输出文档
        if not output_path:
            path = Path(file_path)
            output_path = str(path.parent / f"{path.stem}_zh{path.suffix}")
        
        print("6. 保存文档...")
        self.save_document(final, structure, file_path, output_path)
        
        # 显示报告
        print("\n" + "=" * 60)
        print("✅ 翻译完成！")
        print("=" * 60)
        print(f"输出：{output_path}")
        print(f"总字数：{total_chars:,} 字")
        print(f"用时：{elapsed:.1f} 秒")
        print(f"速度：{total_chars/elapsed:.0f} 字/秒")
        
        final_cost = checkpoint.load().get('total_cost', 0)
        print(f"估算成本：¥{final_cost:.2f}")
        
        # 清理缓存
        checkpoint.clear()
        
        return output_path
    
    async def translate_batch_files(self, input_dir: str, output_dir: Optional[str] = None, max_concurrent: int = 3):
        """批量翻译目录"""
        input_path = Path(input_dir)
        if not output_dir:
            output_path = input_path.parent / f"{input_path.name}_zh"
        else:
            output_path = Path(output_dir)
        
        output_path.mkdir(exist_ok=True)
        
        # 扫描文件
        files = list(input_path.glob("*.docx")) + list(input_path.glob("*.md")) + list(input_path.glob("*.txt"))
        
        if not files:
            print(f"目录 {input_dir} 中没有可翻译的文件")
            return
        
        print(f"\n发现 {len(files)} 个文件")
        print("=" * 60)
        
        total_cost = 0
        for i, file in enumerate(files, 1):
            print(f"\n文件 {i}/{len(files)}: {file.name}")
            output_file = output_path / f"{file.stem}_zh{file.suffix}"
            
            try:
                result = await self.translate_document(str(file), str(output_file), max_concurrent)
                print(f"✅ 完成：{result}")
            except Exception as e:
                print(f"❌ 失败：{e}")
        
        print("\n" + "=" * 60)
        print(f"批量翻译完成！共 {len(files)} 个文件")
        print(f"输出目录：{output_path}")


async def main():
    """命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(description="Smart Translator - 智能翻译官")
    parser.add_argument("input", help="输入文件或目录")
    parser.add_argument("-o", "--output", help="输出文件或目录")
    parser.add_argument("-m", "--model", default="qwen-plus", help="大模型 (qwen-plus, gpt-4o, claude-3.5-sonnet)")
    parser.add_argument("--mode", default="technical", help="翻译模式 (strict, technical, fluent)")
    parser.add_argument("--batch", action="store_true", help="批量翻译目录")
    parser.add_argument("-c", "--concurrent", type=int, default=5, help="最大并发数")
    
    args = parser.parse_args()
    
    # 创建翻译器
    translator = SmartTranslator(model=args.model, mode=args.mode)
    
    if args.batch:
        # 批量翻译
        await translator.translate_batch_files(args.input, args.output, args.concurrent)
    else:
        # 单个文件
        await translator.translate_document(args.input, args.output, args.concurrent)


if __name__ == "__main__":
    asyncio.run(main())
